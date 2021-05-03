from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
import config
import datetime
import calendar
import subprocess
import webbrowser
import os
import shutil
import csv


def get_invoice_date_from_user_rehasport():
    months = ["styczeń", "luty", "marzec", "kwiecień", "maj", "czerwiec", "lipiec", "sierpień", "wrzesień",
              "październik", "listopad", "grudzień"]

    while True:
        invoice_month_word = input("Podaj miesiąc w którym wykonana była usługa:")
        invoice_month_word = invoice_month_word.lower()
        if invoice_month_word in months:
            break
        else:
            print("Nie ma takiego miesiąca spróbuj jeszcze raz:\n")
    number_of_month = months.index(invoice_month_word)+1
    today = datetime.date.today()
    current_year = today.year
    number_of_last_day = calendar.monthrange(current_year, number_of_month)[1]
    if number_of_month == 12:
        current_year = int(today.year) - 1
        check_year = input("Czy rok się zgadza? y/n: " + str(current_year))
        check_year = check_year.lower()
        if check_year != "y" or check_year != "yes":
            current_year = input("Wprowadź rok manualnie: ")
    date_of_invoice = today.replace(day=number_of_last_day, month=int(number_of_month), year=int(current_year))
    return invoice_month_word, date_of_invoice, current_year, number_of_month


def get_invoice_date_from_user_manual():
    pass


def generate_date_of_payment(date_of_invoice):
    date_of_payment = date_of_invoice + datetime.timedelta(days=14)
    return date_of_payment


def get_invoice_amount_from_user():
    while True:
        amount = input("Podaj kwotę pieniędzy na jaką ma być wystawiona faktura: ")
        amount = amount.replace(",", ".")
        if int(amount):
            amount = amount + ".00"
        try:
            float(amount)
        except ValueError:
            print("Podano błędną wartość, spróbuj jeszcze raz! \n")
            continue
        break
    return amount


def generate_full_file_name(invoice_month_word, current_year):
    full_file_name = config.file_name + " - " + invoice_month_word + " " + str(current_year)
    return full_file_name


def generate_invoice_document(invoice_month_word, invoice_amount, date_of_invoice, date_of_payment, full_file_name):
    def find_and_replace(find, replace):
        for paragraph in invoice_doc.paragraphs:
            if find in paragraph.text:
                paragraph.add_run(replace).bold = True

    invoice_doc = Document('template_inv.docx')
    # HEADER
    nr_of_invoice = "1"

    find_and_replace("Faktura nr FV", date_of_invoice.strftime(f"{nr_of_invoice}/%m/%Y"))
    find_and_replace("Data wystawienia:", date_of_invoice.strftime("%d/%m/%Y"))
    find_and_replace("Miejsce wystawienia:", config.invoice_city)

    # TABLES
    invoice_doc.tables[0].cell(1, 0).text = config.company_name.upper()
    invoice_doc.tables[0].cell(2, 0).text = config.company_address_part1
    invoice_doc.tables[0].cell(3, 0).text = config.company_address_part2
    invoice_doc.tables[0].cell(4, 0).text = config.company_nip
    invoice_doc.tables[0].cell(5, 0).text = config.company_email
    invoice_doc.tables[0].cell(6, 0).text = config.company_tel
    invoice_doc.tables[0].cell(1, 0).paragraphs[0].runs[0].font.bold = True

    invoice_doc.tables[0].cell(1, 1).text = config.buyer_name.upper()
    invoice_doc.tables[0].cell(2, 1).text = config.buyer_address_part1
    invoice_doc.tables[0].cell(3, 1).text = config.buyer_address_part2
    invoice_doc.tables[0].cell(4, 1).text = config.buyer_nip
    invoice_doc.tables[0].cell(1, 1).paragraphs[0].runs[0].font.bold = True

    invoice_doc.tables[1].cell(1, 1).text = f"Usługi medyczne wykonane w miesiącu - {invoice_month_word} " \
                                            f"{date_of_invoice.strftime('%Y')}"
    invoice_doc.tables[1].cell(1, 5).text = invoice_amount
    invoice_doc.tables[1].cell(1, 7).text = invoice_amount
    invoice_doc.tables[1].cell(1, 9).text = invoice_amount
    invoice_doc.tables[1].cell(1, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    invoice_doc.tables[1].cell(1, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    invoice_doc.tables[1].cell(1, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    invoice_doc.tables[2].cell(1, 1).text = invoice_amount
    invoice_doc.tables[2].cell(1, 3).text = invoice_amount
    invoice_doc.tables[2].cell(2, 1).text = invoice_amount
    invoice_doc.tables[2].cell(2, 3).text = invoice_amount
    invoice_doc.tables[2].cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    invoice_doc.tables[2].cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    invoice_doc.tables[2].cell(2, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    invoice_doc.tables[2].cell(2, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    invoice_doc.tables[3].cell(0, 3).text = invoice_amount + " PLN"
    invoice_doc.tables[3].cell(2, 3).text = invoice_amount + " PLN"
    invoice_doc.tables[3].cell(2, 1).text = config.bank_name
    invoice_doc.tables[3].cell(3, 1).text = config.bank_account
    invoice_doc.tables[3].cell(3, 3).text = "word_value"
    invoice_doc.tables[3].cell(1, 1).text = date_of_payment.strftime("%d/%m/%Y")

    invoice_doc.tables[4].cell(0, 0).text = config.name_of_issuing
    invoice_doc.tables[4].cell(0, 0).paragraphs[0].runs[0].font.bold = True
    invoice_doc.tables[4].cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    invoice_doc.save(full_file_name + ".docx")


def convert_to_pdf(full_file_name):
    if config.os_user == "1":
        pass
    elif config.os_user == "2":
        subprocess.check_output(['libreoffice', '--convert-to', 'pdf', full_file_name + ".docx"])
    else:
        print("Podano złą wartość w pliku config 'os_user'")


def open_preview(full_file_name):
    webbrowser.open(full_file_name + ".pdf", new=1)


def confirm_invoice_from_user(invoice_month_word, full_file_name):
    while True:
        approval = input("Czy faktura jest ok? y/n: ")
        approval = approval.lower()
        if approval == "y":
            os.makedirs(os.path.dirname(config.path + invoice_month_word + "/"), exist_ok=True)
            shutil.move(full_file_name + ".docx", config.path + invoice_month_word + "/" + full_file_name + ".docx")
            shutil.move(full_file_name + ".pdf", config.path + invoice_month_word + "/" + full_file_name + ".pdf")
            break
        elif approval == "n":
            while True:
                x = input("\nWybierz opcję: \ndel - usuń plik \nedit - zostaw plik .docx: ")
                if x == "del":
                    os.remove(full_file_name + ".docx")
                    os.remove(full_file_name + ".pdf")
                    print("\nUsunięto wszystkie pliki")
                    break
                if x == "edit":
                    os.makedirs(os.path.dirname(config.path + invoice_month_word + "/"), exist_ok=True)
                    shutil.move(full_file_name + ".docx", config.path + invoice_month_word + full_file_name + ".docx")
                    os.remove(full_file_name + ".pdf")
                    print("\nPopraw błędy w pliku .docx \n Plik znajduje się w katalogu miesiąca")
                    break
                else:
                    continue
            break
        else:
            continue


def main():
    invoice_month_word, date_of_invoice, current_year, number_of_month = get_invoice_date_from_user_rehasport()
    date_of_payment = generate_date_of_payment(date_of_invoice)
    invoice_amount = get_invoice_amount_from_user()
    full_file_name = generate_full_file_name(invoice_month_word, current_year)

    print(f"\nMIESIĄC WYSTAWIENIA FAKTURY: {invoice_month_word}")
    print(f"KWOTA NA FAKTURZE: {invoice_amount} PLN")
    generate_invoice_document(invoice_month_word, invoice_amount, date_of_invoice, date_of_payment, full_file_name)
    convert_to_pdf(full_file_name)
    open_preview(full_file_name)
    confirm_invoice_from_user(invoice_month_word, full_file_name)


if __name__ == '__main__':
    main()
